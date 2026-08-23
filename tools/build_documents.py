from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

BLUE = "2E74B5"
DARK = "203748"
DEEP = "0B4F4A"
TEAL = "176F65"
GOLD = "D99B21"
LIGHT = "E8EEF5"
MINT = "EAF3ED"
CREAM = "FFF7DF"
MUTED = "596A70"
WHITE = "FFFFFF"
RED = "B35C54"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure(doc, short_title):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    # Keep a generous, consistent breathing space between the running header
    # rule and the first title or paragraph on every page.
    section.top_margin = Cm(2.35)
    section.header_distance = Cm(0.8)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Title", 30, DARK, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(.375)
        style.paragraph_format.first_line_indent = Inches(-.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = short_title
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    ppr = header._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "D7DEDA")
    border.append(bottom)
    ppr.append(border)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Seite ")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(footer)

    props = doc.core_properties
    props.author = "Lernmaterial für David"
    props.subject = "NMG Weltreligionen – Prüfung Teil 1"
    props.keywords = "NMG, Weltreligionen, 6. Klasse, Zug"


def paragraph(doc, text="", bold=False, italic=False, color=None, size=None, align=None, after=None, keep_together=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if keep_together:
        p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return p


def cover(doc, kicker, title, subtitle, audience, language_line, package_line="—  Lernpaket zur ersten Prüfung  —"):
    paragraph(doc, "", after=76)
    paragraph(doc, kicker.upper(), bold=True, color=GOLD, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    paragraph(doc, title, bold=True, color=DARK, size=30, align=WD_ALIGN_PARAGRAPH.CENTER, after=9)
    paragraph(doc, subtitle, color="2B5163", size=15, align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
    paragraph(doc, package_line, color=GOLD, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=72)
    paragraph(doc, "DAVID · 6. KLASSE · KANTON ZUG", bold=True, color=DARK, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
    paragraph(doc, audience, italic=True, color=MUTED, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=17)
    paragraph(doc, language_line, color=TEAL, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc, items):
    num_id = fresh_numbering_id(doc, 1)
    for item in items:
        p = doc.add_paragraph(item)
        apply_numbering(p, num_id)


def fresh_numbering_id(doc, start=1):
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_num = style.find(".//" + qn("w:numId"))
    base_num_id = style_num.get(qn("w:val"))
    base_num = next(node for node in numbering.findall(qn("w:num")) if node.get(qn("w:numId")) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_node = OxmlElement("w:startOverride")
    start_node.set(qn("w:val"), str(start))
    override.append(start_node)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)


def callout(doc, title, text, fill=MINT, accent=TEAL):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(.34)
    table.columns[1].width = Inches(6.16)
    left, right = table.rows[0].cells
    shade(left, accent)
    shade(right, fill)
    for cell in (left, right):
        set_cell_margins(cell, 100, 100, 130, 130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left.text = "✓"
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    left.paragraphs[0].runs[0].font.bold = True
    p = right.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DEEP)
    p.add_run(text)
    paragraph(doc, "", after=2)


def table(doc, headers, rows, widths=None, font_size=9.2):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    tbl.autofit = False
    if widths:
        for i, width in enumerate(widths):
            tbl.columns[i].width = Inches(width)
    header = tbl.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        shade(cell, LIGHT)
        set_cell_margins(cell)
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(DEEP)
    for row_data in rows:
        row = tbl.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell)
            cell.text = str(text)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    run.font.size = Pt(font_size)
    paragraph(doc, "", after=1)
    return tbl


def label_table(doc, rows):
    return table(doc, ["Begriff", "Kurz erklärt"], rows, [1.18, 5.32], 9.6)


def qblock(doc, questions, answers=False, start=1):
    num_id = fresh_numbering_id(doc, start) if not answers else None
    for i, item in enumerate(questions, 1):
        if answers:
            q, a = item
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            r = p.add_run(f"{i}. {q}\n")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(DEEP)
            p.add_run(a)
        else:
            p = doc.add_paragraph()
            apply_numbering(p, num_id)
            p.add_run(item)
            for _ in range(2):
                blank = doc.add_paragraph("________________________________________________________________________")
                blank.paragraph_format.space_after = Pt(2)
                blank.runs[0].font.color.rgb = RGBColor.from_string("CCD3D0")


WORLD_ROWS_DE = [
    ["Christentum", "Kreuz", "weltweit; u. a. Europa/Amerika", "ein Gott; Dreifaltigkeit", "Kirche", "Bibel"],
    ["Islam", "Halbmond*", "weltweit; u. a. Nordafrika/West- und Südasien", "ein Gott (Allah)", "Moschee", "Koran"],
    ["Judentum", "Davidstern", "weltweit; besonders Israel/USA", "ein Gott", "Synagoge", "Tora/Tanach"],
    ["Buddhismus", "Dharma-Rad", "Ost- und Südostasien", "kein Schöpfergott im Zentrum", "Tempel", "verschiedene Lehrtexte"],
    ["Hinduismus", "Om", "vor allem Indien/Südasien", "vielfältige Gottesvorstellungen", "Tempel", "z. B. Veden"],
]


def build_german():
    doc = Document()
    configure(doc, "Davids Weltreligionen-Training · Lernheft DE")
    cover(doc, "NMG · Prüfung Teil 1", "Weltreligionen verstehen", "Lernheft mit Prüfungswissen, Merksätzen und Probeprüfung", "Für David und seine Lernbegleitung", "Deutsch · Schweizer Schreibweise")

    h1(doc, "1. Was gehört zur ersten Prüfung?")
    callout(doc, "Der genaue Umfang", "Gemeinsame Ziele zu allen fünf Weltreligionen sowie die Detailziele zu Christentum, Islam und Judentum. Buddhismus und Hinduismus werden in Teil 1 nur im Überblick geprüft.")
    h2(doc, "Ich kann …")
    bullets(doc, [
        "die fünf Weltreligionen nennen und ihre Verbreitung grob beschreiben.",
        "Monotheismus und Polytheismus erklären und Beispiele zuordnen.",
        "Symbole, Gotteshäuser und heilige Schriften erkennen.",
        "Christentum, Islam und Judentum nach denselben Merkmalen vergleichen.",
        "wichtige Stationen aus dem Leben Jesu und Mohammeds erklären.",
        "eine katholische Kirche und eine Moschee mit Fachbegriffen beschriften.",
    ])
    h2(doc, "So lernst du wirksam")
    numbers(doc, ["Lies eine Karte laut und decke sie ab.", "Erkläre sie ohne abzulesen in zwei Sätzen.", "Trainiere 10 Fragen in der App.", "Wiederhole am nächsten Tag zuerst die schwächsten Themen."])
    callout(doc, "5-Tage-Plan", "Tag 1 Überblick · Tag 2 Christentum/Jesus · Tag 3 Islam/Mohammed · Tag 4 Judentum und Gebäude · Tag 5 Probeprüfung und Fehler wiederholen.", CREAM, GOLD)

    page_break(doc)
    h1(doc, "2. Überblick: fünf Weltreligionen")
    paragraph(doc, "Religionen sind Weltanschauungen und gelebte Traditionen. Sie beschäftigen sich mit Sinnfragen, Gemeinschaft, Werten, Festen und Vorstellungen über Gott oder das Heilige. Nicht jede Religion beantwortet diese Fragen gleich.")
    table(doc, ["Religion", "Symbol", "Verbreitung", "Gottesvorstellung", "Gebäude", "Schrift"], WORLD_ROWS_DE, [1.0, .75, 1.35, 1.25, .95, 1.2], 8.2)
    paragraph(doc, "* Für die Prüfung gilt die Symbolzuordnung des Lernhefts. Der Halbmond ist verbreitet, aber kein überall verbindliches offizielles Symbol des Islams.", italic=True, color=MUTED, size=8.5)
    h2(doc, "Mono oder poly?")
    label_table(doc, [["Monotheismus", "Glaube an einen Gott. Beispiele: Judentum, Christentum, Islam."], ["Polytheismus", "Glaube an mehrere Gottheiten. In hinduistischen Traditionen gibt es vielfältige Vorstellungen; eine einfache Zuordnung kann zu kurz greifen."], ["Buddhismus", "Ein Schöpfergott steht nicht im Zentrum. Deshalb passt weder „mono“ noch „poly“ immer gut."]])
    callout(doc, "Merksatz", "MONO = eins · POLY = viele. Karten zeigen Schwerpunkte, keine Grenzen: Religionen sind heute weltweit verbreitet.")

    page_break(doc)
    h1(doc, "3. Christentum und Jesus")
    label_table(doc, [["Symbol", "Kreuz"], ["Gott", "ein Gott; Vater, Sohn und Heiliger Geist (Dreifaltigkeit)"], ["Gotteshaus", "Kirche"], ["Heilige Schrift", "Bibel: Altes und Neues Testament"], ["Konfessionen", "katholisch, reformiert/evangelisch, orthodox"], ["Regeln/Bräuche", "Nächstenliebe, Zehn Gebote, Gebet, Gottesdienst, Weihnachten, Karfreitag, Ostern, Himmelfahrt"]])
    h2(doc, "Jesus – vier Stationen")
    table(doc, ["Zeit", "Station", "Warum wichtig?"], [
        ["vor über 2000 Jahren", "Jude aus Nazareth; nach den Evangelien Geburt in Bethlehem", "Jesus gehört geschichtlich zum Judentum."],
        ["als Erwachsener", "Wanderprediger; spricht von Gottes Liebe und Nächstenliebe", "Seine Botschaft gibt Menschen Hoffnung."],
        ["ca. 30/33 n. Chr.", "Kreuzigung in Jerusalem unter römischer Herrschaft", "Karfreitag erinnert daran."],
        ["christlicher Glaube", "Auferstehung am dritten Tag; später Himmelfahrt", "Ostern ist das wichtigste christliche Fest."],
    ], [1.15, 2.7, 2.65], 9.2)
    callout(doc, "Sauber formulieren", "Geschichte: „Jesus war Jude und wurde gekreuzigt.“ Glaubensaussage: „Christinnen und Christen glauben, dass Jesus der Sohn Gottes ist und auferstand.“", CREAM, GOLD)
    h2(doc, "Innenraum einer katholischen Kirche")
    table(doc, ["Begriff", "Funktion"], [["Altar", "Tisch der Eucharistiefeier"], ["Ambo", "Lesepult für Bibeltexte und Predigt"], ["Tabernakel", "Aufbewahrungsort für geweihte Hostien"], ["Taufbecken", "Ort der Taufe"], ["Kruzifix", "Kreuz mit Darstellung Jesu"], ["Orgel", "Instrument für Gottesdienst und Gesang"]], [1.35, 5.15], 9.5)

    h1(doc, "4. Islam und Mohammed")
    label_table(doc, [["Symbol", "Halbmond im Lernheft"], ["Gott", "Allah – das arabische Wort für Gott; ein Gott"], ["Gotteshaus", "Moschee"], ["Heilige Schrift", "Koran"], ["Wichtige Person", "Mohammed, im Islam Prophet und Gesandter Gottes"], ["Regeln/Bräuche", "fünf Säulen, Ramadan, Feste, Gebet und Gemeinschaft"]])
    h2(doc, "Mohammed – vier Stationen")
    table(doc, ["Zeit", "Station", "Prüfungswort"], [["ca. 570", "Geburt in Mekka; früh Waise, später Händler", "Arabien · Onkel · Händler"], ["ca. 610", "nach islamischer Überlieferung erste Offenbarung durch Gabriel in der Höhle Hira", "Höhle bei Mekka"], ["622", "Auswanderung von Mekka nach Medina", "Hidschra · Medina"], ["632", "Tod in Medina; Offenbarungen werden später im Koran gesammelt", "Prophet · Koran"]], [1.1, 3.6, 1.8], 9.2)
    callout(doc, "Achtung: Druckfehler", "Eine Aufgabe fragt nach einer Höhle „bei Medina“. Gemeint ist die Höhle Hira bei Mekka. Für das Kreuzworträtsel lautet das gesuchte Wort trotzdem HÖHLE.", CREAM, GOLD)
    h2(doc, "Die fünf Säulen")
    table(doc, ["Säule", "Kurz erklärt"], [["Glaubensbekenntnis", "Bekenntnis zu dem einen Gott und Mohammed als seinem Gesandten"], ["Gebet", "regelmässiges Gebet"], ["Almosen", "Teilen und Bedürftige unterstützen"], ["Fasten", "Fastenmonat Ramadan"], ["Pilgerfahrt", "wenn möglich Pilgerfahrt nach Mekka"]], [1.7, 4.8], 9.5)
    h2(doc, "Eine Moschee beschriften")
    label_table(doc, [["Mihrab", "Gebetsnische; zeigt die Qibla, die Richtung nach Mekka"], ["Minbar", "Kanzel, besonders für die Predigt am Freitag"], ["Gebetsraum", "offener Raum für das gemeinsame Gebet"], ["Gebetsteppiche", "markieren saubere Gebetsplätze und Reihen"], ["Waschung", "rituelle Reinigung vor dem Gebet"], ["Minarett/Kuppel", "häufige, aber nicht zwingende äussere Merkmale"]])

    page_break(doc)
    h1(doc, "5. Judentum")
    label_table(doc, [["Symbol", "Davidstern"], ["Gott", "ein Gott"], ["Gotteshaus", "Synagoge – Ort für Gebet, Lernen und Gemeinschaft"], ["Heilige Schriften", "Tora (fünf Bücher Mose) als Kern; Tanach"], ["Wichtige Personen", "z. B. Abraham und Mose"], ["Regeln/Bräuche", "Gebote, Schabbat, Feste, unterschiedliche Formen koscherer Lebensweise"]])
    h2(doc, "In einer Synagoge")
    table(doc, ["Begriff", "Funktion"], [["Toraschrein", "Aufbewahrungsort für Torarollen"], ["Bima", "erhöhtes Lesepult für die Toralesung"], ["Ewiges Licht", "Licht nahe dem Toraschrein"], ["Torarolle", "handgeschriebener Text der Tora"], ["Kippa", "Kopfbedeckung, die viele jüdische Männer und Jungen tragen"]], [1.35, 5.15], 9.5)
    h2(doc, "Gemeinsamkeiten der drei Religionen")
    bullets(doc, ["Glaube an einen Gott.", "Verbindung zu Abraham und biblischen Überlieferungen.", "heilige Schriften, Gebete, Feste und Regeln.", "Jerusalem hat in allen drei Traditionen Bedeutung.", "Gemeinschaft und Hilfe für andere spielen eine wichtige Rolle."])
    callout(doc, "Prüfungs-Merksatz", "Kirche–Bibel–Jesus · Moschee–Koran–Mohammed · Synagoge–Tora–Mose")

    page_break(doc)
    h1(doc, "6. Probeprüfung")
    paragraph(doc, "Arbeite ohne Lernheft. Schreibe kurze, klare Sätze. Kontrolliere danach mit der App oder der Lösungsbroschüre.", italic=True, color=MUTED)
    qblock(doc, [
        "Nenne die fünf Weltreligionen.",
        "Erkläre Monotheismus und nenne drei Beispiele.",
        "Erkläre Polytheismus.",
        "Ordne zu: Kirche, Moschee, Synagoge – Bibel, Koran, Tora.",
        "Nenne Symbol, Gotteshaus und heilige Schrift des Christentums.",
        "Nenne drei christliche Konfessionen.",
        "Erzähle drei Stationen aus dem Leben Jesu.",
        "Was feiern Christinnen und Christen an Ostern?",
        "Beschrifte eine katholische Kirche mit vier Begriffen.",
        "Was bedeutet Allah?",
    ])
    page_break(doc)
    h1(doc, "Probeprüfung – Fortsetzung")
    qblock(doc, [
        "Erzähle vier Stationen aus dem Leben Mohammeds.",
        "Wo liegt die Höhle Hira?",
        "Nenne die fünf Säulen des Islams.",
        "Beschrifte eine Moschee mit vier Begriffen.",
        "Nenne Symbol, Gotteshaus und heilige Schrift des Judentums.",
        "Was ist die Tora?",
        "Erkläre den Schabbat.",
        "Nenne drei Gemeinsamkeiten von Judentum, Christentum und Islam.",
        "Warum sind Weltkarten zu Religionen immer Vereinfachungen?",
        "Formuliere eine Glaubensaussage respektvoll.",
    ], start=11)

    page_break(doc)
    h1(doc, "7. Für Eltern: David wirksam begleiten")
    callout(doc, "Ziel", "Nicht möglichst lange lernen, sondern aktiv abrufen: David soll erklären, vergleichen, beschriften und Fehler zeitversetzt wiederholen.")
    h2(doc, "15-Minuten-Routine")
    table(doc, ["Minuten", "Aktivität"], [["0–3", "David nennt ohne Hilfe alles, was er noch weiss."], ["3–8", "Eine Lernkarte gemeinsam klären; David erklärt sie in eigenen Worten."], ["8–13", "5–8 App-Fragen oder mündliche Fragen."], ["13–15", "Ein Fehler wird als Merksatz notiert; mit einem Erfolg abschliessen."]], [1.0, 5.5], 9.6)
    h2(doc, "Gute Rückfragen")
    bullets(doc, ["Woran erkennst du das?", "Was ist ähnlich, was verschieden?", "Ist das eine historische Aussage oder eine Glaubensaussage?", "Kannst du das in zwei Sätzen erklären?", "Welcher Merksatz hilft dir?"])
    h2(doc, "Quellen und Abgleich")
    paragraph(doc, "Fachliche Orientierung: Lehrplan 21 Kanton Zug, besonders NMG.12.1, NMG.12.2 und NMG.12.5. Das Lernzielblatt der Klasse bestimmt den konkreten Prüfungsumfang. Offizielle Seiten: zg.ch/de/bildung/schulen/gemeindliche-schulen/unterricht/lehrplan21 · zg.lehrplan.ch")
    paragraph(doc, "Hinweis: Religiöse Traditionen sind vielfältig. Tabellen und Symbole sind Lernhilfen und keine vollständige Beschreibung aller Menschen einer Religion.", italic=True, color=MUTED, size=9)
    path = OUT / "Lernheft_Pruefung_1_David_DE.docx"
    doc.save(path)
    return path


def build_spanish():
    doc = Document()
    configure(doc, "Material de apoyo ES · NMG Religiones")
    cover(doc, "NMG · Primera prueba", "Religiones del mundo", "Material de estudio en español y guía práctica para la familia", "Para David y sus padres", "Español · términos clave conservados en alemán")
    h1(doc, "1. Alcance exacto de la primera prueba")
    callout(doc, "Qué entra", "Objetivos comunes de las cinco religiones del mundo y, con detalle, cristianismo, islam y judaísmo. Budismo e hinduismo aparecen solamente en el panorama común; sus bloques específicos pertenecen a la segunda prueba.")
    h2(doc, "David debe poder")
    bullets(doc, ["nombrar las cinco religiones y situar sus principales zonas de difusión;", "explicar Monotheismus y Polytheismus;", "comparar símbolo, Dios, edificio, escritura, personas importantes, normas y costumbres;", "contar episodios principales de Jesús y Mahoma;", "identificar partes de una iglesia católica y de una mezquita;", "responder con respeto, distinguiendo historia y afirmaciones de fe."])
    h2(doc, "Glosario alemán–español")
    table(doc, ["Deutsch", "Español", "Idea breve"], [["Gotteshaus", "lugar de culto", "Kirche, Moschee, Synagoge"], ["Heilige Schrift", "escritura sagrada", "Bibel, Koran, Tora"], ["Glaube", "fe/creencia", "convicciones religiosas"], ["Brauch", "costumbre/rito", "práctica vivida"], ["Gebot", "mandamiento", "regla religiosa"], ["Verbreitung", "distribución", "zonas de presencia"], ["Auferstehung", "resurrección", "creencia cristiana sobre Jesús"], ["Offenbarung", "revelación", "mensaje de Dios según una tradición"]], [1.55, 1.85, 3.1], 9.2)

    page_break(doc)
    h1(doc, "2. Panorama de las cinco religiones")
    table(doc, ["Religión", "Símbolo", "Distribución simplificada", "Concepción de Dios", "Lugar", "Texto"], [
        ["Cristianismo", "cruz", "mundial; Europa/América, entre otras", "un Dios; Trinidad", "iglesia", "Biblia"],
        ["Islam", "media luna*", "mundial; norte de África y Asia occidental/meridional", "un Dios (Alá)", "mezquita", "Corán"],
        ["Judaísmo", "estrella de David", "mundial; especialmente Israel/EE. UU.", "un Dios", "sinagoga", "Torá/Tanaj"],
        ["Budismo", "rueda del dharma", "Asia oriental y sudoriental", "sin dios creador central", "templo", "varios textos"],
        ["Hinduismo", "Om", "India y Asia meridional", "concepciones diversas", "templo", "p. ej., Vedas"],
    ], [1.0, .75, 1.45, 1.25, .9, 1.15], 8.1)
    paragraph(doc, "* La media luna es un símbolo cultural muy difundido, pero no es un símbolo oficial universal del islam. Para el examen se sigue la asociación del folleto.", italic=True, color=MUTED, size=8.5)
    h2(doc, "Monoteísmo y politeísmo")
    label_table(doc, [["Monotheismus", "Creencia en un solo Dios: judaísmo, cristianismo e islam."], ["Polytheismus", "Creencia en varias divinidades. Las tradiciones hindúes son diversas y no caben siempre en una sola etiqueta."], ["Budismo", "Un dios creador no está en el centro; mono/poli no es una clasificación suficiente."]])
    callout(doc, "Para memorizar", "MONO = uno · POLY = muchos. Los mapas muestran concentraciones, no fronteras rígidas.")

    page_break(doc)
    h1(doc, "3. Cristianismo y Jesús")
    table(doc, ["Rasgo", "Respuesta para el examen"], [["Symbol", "Kreuz – cruz"], ["Gott", "un Dios; Vater, Sohn und Heiliger Geist – Padre, Hijo y Espíritu Santo"], ["Gotteshaus", "Kirche – iglesia"], ["Heilige Schrift", "Bibel – Biblia"], ["Konfessionen", "katholisch, reformiert/evangelisch, orthodox"], ["Regeln/Bräuche", "amor al prójimo, Diez Mandamientos, Navidad, Viernes Santo, Pascua, Ascensión"]], [1.45, 5.05], 9.5)
    h2(doc, "Jesús en cuatro pasos")
    numbers(doc, ["Fue judío y procedía de Nazaret; según los evangelios nació en Belén.", "Como predicador itinerante habló del amor de Dios y del prójimo.", "Fue crucificado en Jerusalén bajo dominio romano.", "La fe cristiana afirma que resucitó al tercer día y después ascendió al cielo."])
    h2(doc, "Iglesia católica: vocabulario")
    table(doc, ["Alemán", "Español"], [["Altar", "altar, mesa de la eucaristía"], ["Ambo", "ambón/atril de lecturas"], ["Tabernakel", "sagrario"], ["Taufbecken", "pila bautismal"], ["Kruzifix", "crucifijo"], ["Orgel", "órgano"]], [1.7, 4.8], 9.6)
    callout(doc, "Frase modelo", "„Christinnen und Christen glauben, dass Jesus auferstanden ist.“ = Los cristianos creen que Jesús resucitó.", CREAM, GOLD)

    page_break(doc)
    h1(doc, "4. Islam y Mahoma")
    table(doc, ["Rasgo", "Respuesta para el examen"], [["Symbol", "Halbmond – media luna, según el folleto"], ["Gott", "Allah: palabra árabe para Dios; un solo Dios"], ["Gotteshaus", "Moschee – mezquita"], ["Heilige Schrift", "Koran – Corán"], ["Wichtige Person", "Mohammed – Mahoma, profeta y mensajero de Dios en el islam"], ["Regeln/Bräuche", "cinco pilares, Ramadán, oración y fiestas"]], [1.45, 5.05], 9.5)
    h2(doc, "Mahoma en cuatro pasos")
    numbers(doc, ["Nació hacia 570 en La Meca; quedó huérfano y más tarde fue comerciante.", "Según la tradición islámica recibió la primera revelación del ángel Gabriel en la cueva de Hira, cerca de La Meca.", "En 622 emigró de La Meca a Medina: la Hégira.", "Murió en 632 en Medina; las revelaciones fueron recopiladas después en el Corán."])
    callout(doc, "Corrección importante", "El ejercicio dice «cueva cerca de Medina», pero Hira está cerca de La Meca. En el crucigrama la respuesta sigue siendo HÖHLE (cueva).", CREAM, GOLD)
    h2(doc, "Partes de una mezquita")
    table(doc, ["Alemán", "Español"], [["Mihrab", "nicho que indica la dirección de La Meca"], ["Qibla", "dirección de la oración hacia La Meca"], ["Minbar", "púlpito"], ["Gebetsraum", "sala de oración"], ["Gebetsteppich", "alfombra de oración"], ["Waschung", "ablución o lavado ritual"]], [1.7, 4.8], 9.5)

    page_break(doc)
    h1(doc, "5. Judaísmo y comparación")
    table(doc, ["Rasgo", "Respuesta para el examen"], [["Symbol", "Davidstern – estrella de David"], ["Gott", "un solo Dios"], ["Gotteshaus", "Synagoge – sinagoga"], ["Heilige Schrift", "Tora – Torá; núcleo de las escrituras judías"], ["Wichtige Personen", "por ejemplo, Abraham y Moisés"], ["Regeln/Bräuche", "mandamientos, Shabat, fiestas y distintas prácticas kosher"]], [1.45, 5.05], 9.5)
    h2(doc, "Tres parejas que no se pueden confundir")
    callout(doc, "Regla de memoria", "Kirche–Bibel–Jesus · Moschee–Koran–Mohammed · Synagoge–Tora–Mose")
    h2(doc, "Similitudes")
    bullets(doc, ["Las tres son monoteístas.", "Tienen escrituras, oraciones, fiestas, normas y comunidades.", "Comparten figuras y relatos relacionados con Abraham.", "Jerusalén tiene importancia en las tres tradiciones."])
    h2(doc, "Diferencias que David debe explicar")
    table(doc, ["Tema", "Cristianismo", "Islam", "Judaísmo"], [["Persona central del tema", "Jesús, Hijo de Dios según la fe cristiana", "Mahoma, profeta y mensajero", "Moisés y otros personajes bíblicos"], ["Escritura", "Biblia", "Corán", "Torá/Tanaj"], ["Lugar", "iglesia", "mezquita", "sinagoga"], ["Rito/tiempo", "Navidad y Pascua", "Ramadán y cinco pilares", "Shabat y fiestas judías"]], [1.1, 1.8, 1.8, 1.8], 8.9)

    page_break(doc)
    h1(doc, "6. Guía para padres")
    callout(doc, "Método recomendado", "Sesiones de 15 minutos, cinco días. Primero recuperación sin mirar; después corrección; al final 5–8 preguntas en la app. Es mejor repetir en varios días que estudiar todo de una vez.")
    h2(doc, "Plan de cinco días")
    table(doc, ["Día", "Tema", "Comprobación"], [["1", "Cinco religiones, símbolos, mapa, mono/poli", "David completa una tabla sin ayuda."], ["2", "Cristianismo, Jesús e iglesia", "Relata cuatro pasos y etiqueta cuatro elementos."], ["3", "Islam, Mahoma y mezquita", "Relata cuatro pasos y nombra cinco pilares."], ["4", "Judaísmo y comparación", "Explica cinco similitudes/diferencias."], ["5", "Prueba simulada", "Repite únicamente los errores."]], [.55, 3.0, 2.95], 9.2)
    h2(doc, "Preguntas orales útiles")
    bullets(doc, ["¿Cómo lo sabes?", "¿Qué tienen en común?", "¿Qué palabra alemana usarías en el examen?", "¿Es un dato histórico o una afirmación de fe?", "Explícamelo en dos frases, sin leer."])
    h2(doc, "Cómo corregir")
    bullets(doc, ["Primero reconocer lo que sí está bien.", "Corregir una sola idea cada vez.", "Pedir que David vuelva a explicarla con sus palabras.", "Crear un ejemplo o una pareja de memoria.", "Revisar el mismo error al día siguiente."])
    h2(doc, "Uso de la aplicación")
    paragraph(doc, "La aplicación guarda estrellas, rachas y dominio por tema en el dispositivo. Conviene practicar en el mismo navegador. El modo adaptativo da prioridad a los temas con menor porcentaje. Puede instalarse como PWA en macOS y Windows y funciona sin conexión después de la primera carga.")

    page_break(doc)
    h1(doc, "7. Mini examen oral con respuestas")
    qblock(doc, [
        ("¿Qué significa Monotheismus?", "Creencia en un solo Dios. Judaísmo, cristianismo e islam son ejemplos."),
        ("¿Qué tres parejas hay que memorizar?", "Kirche–Bibel–Jesus; Moschee–Koran–Mohammed; Synagoge–Tora–Mose."),
        ("¿Qué celebran los cristianos en Pascua?", "La resurrección de Jesús según la fe cristiana."),
        ("¿Dónde está la cueva de Hira?", "Cerca de La Meca, no de Medina."),
        ("¿Qué indica el mihrab?", "La qibla, dirección de la oración hacia La Meca."),
        ("¿Qué es la Torá?", "Los cinco libros de Moisés y núcleo de las escrituras judías."),
        ("¿Qué es el Shabat?", "Día judío de descanso, desde el viernes por la tarde hasta el sábado por la tarde."),
        ("¿Por qué un mapa religioso es una simplificación?", "Porque hoy las religiones están presentes en muchos países y las comunidades son diversas."),
    ], answers=True)
    h2(doc, "Fuentes oficiales")
    paragraph(doc, "Lehrplan 21 del cantón de Zug: zg.ch/de/bildung/schulen/gemeindliche-schulen/unterricht/lehrplan21. Competencias relacionadas: NMG.12.1, NMG.12.2 y NMG.12.5 en zg.lehrplan.ch. El listado de objetivos de la clase determina el contenido concreto de esta primera prueba.")
    paragraph(doc, "Nota: las religiones son diversas. Las tablas son ayudas escolares y no describen a todas las personas o corrientes de una tradición.", italic=True, color=MUTED, size=9)
    path = OUT / "Material_estudio_y_guia_padres_Prueba_1_ES.docx"
    doc.save(path)
    return path


def build_solutions_german():
    doc = Document()
    configure(doc, "Lösungen · NMG Religionen · Prüfung Teil 1")
    cover(doc, "Lösungen · Prüfung 1", "Weltreligionen – Teil 1", "Selbstständig verständliche Lösungen und Lernantworten", "Für Eltern und Lernbegleitende", "Deutsch · Schweizer Schreibweise")
    h1(doc, "Hinweise")
    callout(doc, "Umfang", "Gelöst werden die Unterrichtsaufgaben, soweit sie für Prüfung Teil 1 oder den gemeinsamen Überblick relevant sind. Jede Überschrift beschreibt die Aufgabe in Worten; alle Kernaussagen stehen direkt in diesem Dokument. Es werden weder Fotodateien noch Bildnummern benötigt. Offene persönliche Fragen haben keine einzige richtige Lösung.")
    bullets(doc, ["Kurze Lösungen genügen meist; vollständige Sätze sind bei Erklärfragen besser.", "Drei sachliche Hinweise sind markiert: Hira liegt bei Mekka; Symbolkarten vereinfachen; die 100er-Grafik ist eine Übung und keine aktuelle Weltstatistik."])

    page_break(doc)
    h1(doc, "1. Lernziele für Prüfung Teil 1")
    h2(doc, "Prüfung Teil 1")
    bullets(doc, ["Gemeinsam: fünf Weltreligionen, grobe Verbreitung, Monotheismus/Polytheismus, Fragen und Unterrichtsinhalte.", "Christentum: Merkmale, Jesus, katholische Kirche.", "Islam: Merkmale, Mohammed, Moschee.", "Judentum: Merkmale, Regeln und Bräuche."])
    callout(doc, "Nicht im Detail für Teil 1", "Die eigenen Detailblöcke zu Buddhismus und Hinduismus sind grün markiert und gehören zu Prüfung Teil 2. Im gemeinsamen Überblick können Namen, Symbole und Verbreitung trotzdem vorkommen.")
    h2(doc, "Musterantwort: Monotheismus / Polytheismus")
    paragraph(doc, "Monotheismus bedeutet Glaube an einen Gott. Judentum, Christentum und Islam sind monotheistisch. Polytheismus bedeutet Glaube an mehrere Gottheiten. Bei Hinduismus und Buddhismus sind einfache Etiketten nur begrenzt passend.")

    h1(doc, "2. Lückentext: Grundwissen zu Weltreligionen")
    paragraph(doc, "Lösungswörter in Reihenfolge:", bold=True, color=DEEP)
    table(doc, ["1–8", "9–15", "16–22"], [["lateinischen\nhöhere Macht\nwissenschaftlich\nHalt\nWeltreligionen\nbesonders stark\neinen\nPolytheismus", "Wertvorstellungen\nCharakter\nVerhalten\nBauten\nKirche\nMoschee\nSynagoge", "Feiertagen\nTod\nunterschiedliche\nfrei wählen\ngar keine\nrespektieren\nbehandeln"]], [2.15, 2.15, 2.2], 9.5)
    page_break(doc)
    h2(doc, "Vollständige Kernaussage ohne Lückentext")
    paragraph(doc, "Das Wort Religion hat einen lateinischen Ursprung. Viele Religionen beziehen sich auf eine höhere Macht oder auf das Heilige; solche Glaubensaussagen lassen sich nicht wie naturwissenschaftliche Aussagen beweisen. Religion kann Menschen Halt geben. Zu den fünf Weltreligionen zählen Christentum, Islam, Judentum, Hinduismus und Buddhismus. Monotheistische Religionen glauben an einen Gott; Polytheismus bezeichnet den Glauben an mehrere Gottheiten. Religionen prägen Wertvorstellungen, Charakter und Verhalten. Sichtbar werden sie auch in Bauten wie Kirche, Moschee und Synagoge sowie in Festen und Vorstellungen über Leben und Tod. Menschen dürfen ihren Glauben frei wählen oder keiner Religion angehören. Andere Überzeugungen sollen respektiert und Menschen fair behandelt werden.", keep_together=True)

    h1(doc, "3. Verständnisfragen zum Grundtext")
    qblock(doc, [
        ("Was sind Religionen?", "Religionen sind Weltanschauungen und gelebte Traditionen. Oft gehört der Glaube an eine höhere Macht oder an das Heilige dazu; Religion prägt Werte, Rituale und Gemeinschaft."),
        ("Wie heissen die fünf Weltreligionen?", "Christentum, Islam, Judentum, Hinduismus und Buddhismus."),
        ("Was beinhaltet der Glaube einer Religion?", "Vorstellungen über Gott oder das Heilige, Werte und Verhalten, Regeln, Bräuche, Feste, Gebet, Gemeinschaft und Fragen nach Leben und Tod."),
        ("Nenne einen heiligen Bau.", "Zum Beispiel Kirche, Moschee, Synagoge oder Tempel."),
    ], answers=True)
    callout(doc, "Korrektur zur Notiz", "Die Herleitung des Wortes „Religion“ ersetzt keine Antwort auf die Frage nach den Inhalten des Glaubens. Gefragt sind Vorstellungen, Werte, Verhalten und Praxis.", CREAM, GOLD)

    h1(doc, "4. Weltkarte: Farben, Symbole und Verbreitung")
    table(doc, ["Religion", "Farbe auf der Karte", "Symbol im Heft", "Schwerpunkt vereinfacht"], [["Judentum", "blau", "Davidstern", "Israel; weltweit"], ["Buddhismus", "orange", "Dharma-Rad", "Ost-/Südostasien"], ["Christentum", "violett", "Kreuz", "Europa, Amerika, Teile Afrikas, Australien"], ["Hinduismus", "gelb", "Om", "Indien/Südasien"], ["Islam", "grün", "Halbmond", "Nordafrika, West- und Südasien"]], [1.15, 1.15, 1.25, 3.0], 9.0)
    paragraph(doc, "Die Karte zeigt Schwerpunkte. Heute leben Angehörige aller fünf Religionen in vielen Weltregionen.", italic=True, color=MUTED, size=9)

    page_break(doc)
    h1(doc, "5. 100er-Grafik zur religiösen Zugehörigkeit")
    table(doc, ["Gruppe im Arbeitsblatt", "Anzahl Felder", "Prozent der Grafik"], [["Christen", "31", "31 %"], ["Muslime", "23", "23 %"], ["Hindus", "15", "15 %"], ["Juden", "1", "1 %"], ["Buddhisten", "7", "7 %"], ["andere/keine Religion", "23", "23 %"], ["Summe", "100", "100 %"]], [2.8, 1.65, 2.05], 9.6)
    callout(doc, "Wichtig", "Das ist die Lösung der gezeichneten 100er-Grafik. Sie ist keine verlässliche aktuelle Weltbevölkerungsstatistik.", CREAM, GOLD)
    h1(doc, "6. Persönliche Sinn- und Glaubensfragen")
    paragraph(doc, "Die persönlichen Antworten sind individuell. Es gibt nicht nur eine richtige Lösung. Eine mögliche Musterantwort:")
    table(doc, ["Frage", "Mögliche persönliche Antwort"], [["Wie entstand die Welt?", "Die Naturwissenschaft erklärt Entwicklungen; Religionen erzählen zusätzlich Sinn- und Ursprungsgeschichten."], ["Was ist der Sinn des Lebens?", "Für mich: Beziehungen pflegen, lernen, helfen und Verantwortung übernehmen."], ["Gibt es Gott?", "Menschen antworten verschieden. Ich kann Überzeugungen respektieren, auch wenn ich sie nicht teile."], ["Wie sollen Menschen zusammenleben?", "Fair, friedlich, hilfsbereit und mit Respekt."], ["Was kommt nach dem Tod?", "Religionen geben unterschiedliche Antworten; sicher wissen können wir es nicht."]], [2.0, 4.5], 9.2)

    page_break(doc)
    h1(doc, "7. Vergleichstabelle der fünf Weltreligionen")
    table(doc, ["Religion", "Verbreitung", "Alter/Zeit", "Gott / Personen", "Haus", "Schrift", "Regeln/Bräuche"], [
        ["Christentum", "weltweit", "ca. 2000 J.", "ein Gott; Jesus", "Kirche", "Bibel", "Nächstenliebe, Feste"],
        ["Islam", "weltweit", "ca. 1400 J.", "Allah; Mohammed", "Moschee", "Koran", "fünf Säulen, Ramadan"],
        ["Judentum", "weltweit", "über 3000 J.", "ein Gott; Abraham/Mose", "Synagoge", "Tora/Tanach", "Gebote, Schabbat"],
        ["Buddhismus", "Asien; weltweit", "ca. 2500 J.", "Buddha; kein Schöpfergott zentral", "Tempel", "Lehrtexte", "achtfacher Pfad"],
        ["Hinduismus", "Indien; weltweit", "über 3000 J.", "vielfältige Gottesvorstellungen", "Tempel", "u. a. Veden", "Puja, Feste"],
    ], [.85, .85, .75, 1.25, .7, .85, 1.25], 7.7)
    paragraph(doc, "Altersangaben sind gerundet und dienen nur der Orientierung.", italic=True, color=MUTED, size=8.5)

    h1(doc, "8. Aufgaben zu Jesus und dem Christentum")
    table(doc, ["Aufgabe", "Kurzlösung"], [["Wer war Jesus?", "Jude aus Nazareth; Prediger. Im christlichen Glauben Sohn Gottes/Gott in menschlicher Gestalt."], ["Botschaft", "Gottes Liebe und Nächstenliebe: „Liebe deinen Nächsten wie dich selbst.“"], ["Tod", "Kreuzigung in Jerusalem unter römischer Herrschaft."], ["Danach", "Christlicher Glaube: Auferstehung am dritten Tag; Erscheinungen; Himmelfahrt."], ["Feste", "Weihnachten, Karfreitag, Ostern, Himmelfahrt."]], [1.4, 5.1], 9.4)
    callout(doc, "Präzise Formulierung", "Historisch: Jesus stammte aus Nazareth. Nach den Evangelien wurde er in Bethlehem geboren. Das Blatt formuliert vereinfacht; für die Prüfung die Unterrichtsformulierung beachten.", CREAM, GOLD)

    page_break(doc)
    h1(doc, "9. Aufgaben zu Allah, Mohammed und dem Islam")
    h2(doc, "Allah")
    paragraph(doc, "Allah ist das arabische Wort für Gott. Im Islam gibt es nur einen Gott. Die 99 Namen beschreiben Eigenschaften Gottes, zum Beispiel der Barmherzige oder der Wissende.")
    h2(doc, "Mohammed – Lückenantworten / Kreuzworträtsel")
    table(doc, ["Nr.", "Lösung", "Bezug"], [["1", "Gepriesene", "Bedeutung des Namens im Arbeitsblatt"], ["2", "Arabien", "Herkunft"], ["3", "Onkel", "Erziehung nach dem Tod des Grossvaters"], ["4", "Predigten", "christliche Predigten"], ["5", "Höhle", "Hira bei Mekka"], ["6", "Medina", "Ort der Auswanderung"], ["7", "Tod", "nach dem Tod seiner Frau"], ["Lösungswort", "PROPHET", "Gesandter, der Gottes Botschaft weitergibt"]], [.8, 1.75, 3.95], 9.4)
    callout(doc, "Fehler im Fragetext", "Die Höhle Hira liegt bei Mekka. Wenn in der Aufgabe „bei Medina“ steht, ist das sehr wahrscheinlich ein Druckfehler.", CREAM, GOLD)
    h2(doc, "Musterantwort: Was ist ein Prophet?")
    paragraph(doc, "Ein Prophet ist ein Mensch, der nach religiöser Überlieferung Gottes Botschaft empfängt und weitergibt. Im Islam gilt Mohammed als letzter Prophet.")

    page_break(doc)
    h1(doc, "10. Beschriftungen für Teil 1")
    h2(doc, "Katholische Kirche")
    label_table(doc, [["Altar", "Tisch der Eucharistie"], ["Ambo", "Lesepult"], ["Tabernakel", "Aufbewahrung geweihter Hostien"], ["Taufbecken", "Ort der Taufe"], ["Kruzifix", "Kreuz mit Jesusdarstellung"], ["Orgel", "Instrument für Gottesdienst und Gesang"], ["Weihwasser", "Wasser am Eingang; Erinnerung an die Taufe"]])
    h2(doc, "Moschee")
    label_table(doc, [["Mihrab", "Gebetsnische; zeigt Qibla"], ["Qibla", "Gebetsrichtung nach Mekka"], ["Minbar", "Kanzel"], ["Gebetsraum", "Raum für gemeinsames Gebet"], ["Gebetsteppiche", "Gebetsplätze/Reihen"], ["Waschung", "rituelle Reinigung"], ["Minarett", "Turm; nicht jede Moschee hat eines"]])
    h2(doc, "Synagoge – hilfreich für die Übersicht")
    label_table(doc, [["Toraschrein", "Aufbewahrung der Torarollen"], ["Bima", "Lesepult"], ["Ewiges Licht", "Licht nahe dem Toraschrein"]])

    page_break(doc)
    h1(doc, "11. Schnellkontrolle für Eltern")
    qblock(doc, [
        ("Welche fünf Religionen?", "Christentum, Islam, Judentum, Hinduismus, Buddhismus."),
        ("Welche drei im Detail?", "Christentum, Islam, Judentum."),
        ("Welche drei monotheistisch?", "Judentum, Christentum, Islam."),
        ("Christentum in drei Wörtern?", "Kirche – Bibel – Jesus."),
        ("Islam in drei Wörtern?", "Moschee – Koran – Mohammed."),
        ("Judentum in drei Wörtern?", "Synagoge – Tora – Mose."),
        ("Hira?", "Höhle bei Mekka."),
        ("Ostern?", "Auferstehung Jesu nach christlichem Glauben."),
    ], answers=True)
    h2(doc, "Quellen")
    paragraph(doc, "Grundlage: die bereitgestellten Klassenmaterialien. Fachlicher Abgleich: Lehrplan 21 Kanton Zug, NMG.12.1, NMG.12.2 und NMG.12.5 (zg.lehrplan.ch). Die Lösungen folgen dem Prüfungswortschatz des Hefts und markieren sachliche Vereinfachungen.")
    path = OUT / "Loesungen_Pruefung_1_DE.docx"
    doc.save(path)
    return path


def build_solutions_spanish():
    doc = Document()
    configure(doc, "Soluciones · NMG Religiones · Prueba 1")
    cover(doc, "Soluciones · Prueba 1", "Religiones del mundo - Parte 1", "Soluciones completas y explicaciones autosuficientes", "Para familias y acompañantes del aprendizaje", "Español · términos del examen conservados en alemán", "- Solucionario de la primera prueba -")
    h1(doc, "Indicaciones")
    callout(doc, "Alcance", "Se resuelven las tareas de clase relevantes para la primera prueba o para el panorama común. Cada título describe la actividad con palabras y toda la información esencial aparece en este documento. No se necesitan fotografías ni nombres de archivos. Las preguntas personales abiertas no tienen una única respuesta correcta.")
    bullets(doc, ["Para respuestas explicativas conviene escribir frases completas.", "Se señalan tres precisiones: Hira está cerca de La Meca; los símbolos simplifican; el gráfico de cien casillas es un ejercicio y no una estadística mundial actual."])

    page_break(doc)
    h1(doc, "1. Objetivos de la primera prueba")
    h2(doc, "Contenido de la prueba 1")
    bullets(doc, ["Común: cinco religiones del mundo, distribución general, Monotheismus/Polytheismus, preguntas y contenidos de clase.", "Cristianismo: características, Jesús e iglesia católica.", "Islam: características, Mahoma y mezquita.", "Judaísmo: características, normas y costumbres."])
    callout(doc, "No se estudia en detalle", "Los bloques específicos de budismo e hinduismo pertenecen a la segunda prueba. En el panorama común sí pueden aparecer sus nombres, símbolos y distribución.")
    h2(doc, "Respuesta modelo: monoteísmo y politeísmo")
    paragraph(doc, "Monotheismus significa creer en un solo Dios. Judaísmo, cristianismo e islam son monoteístas. Polytheismus significa creer en varias divinidades. En hinduismo y budismo las etiquetas simples solo resultan parcialmente adecuadas.")

    h1(doc, "2. Texto con huecos: conocimientos básicos")
    paragraph(doc, "Palabras alemanas en el orden de los huecos:", bold=True, color=DEEP)
    table(doc, ["1-8", "9-15", "16-22"], [["lateinischen\nhöhere Macht\nwissenschaftlich\nHalt\nWeltreligionen\nbesonders stark\neinen\nPolytheismus", "Wertvorstellungen\nCharakter\nVerhalten\nBauten\nKirche\nMoschee\nSynagoge", "Feiertagen\nTod\nunterschiedliche\nfrei wählen\ngar keine\nrespektieren\nbehandeln"]], [2.15, 2.15, 2.2], 9.5)
    page_break(doc)
    h2(doc, "Idea completa sin huecos")
    paragraph(doc, "La palabra religión tiene un origen latino. Muchas religiones se refieren a un poder superior o a lo sagrado; esas afirmaciones de fe no se demuestran como una afirmación científica. La religión puede dar apoyo a las personas. Entre las cinco religiones del mundo se cuentan cristianismo, islam, judaísmo, hinduismo y budismo. Las religiones monoteístas creen en un Dios; el politeísmo designa la creencia en varias divinidades. Las religiones influyen en valores, carácter y comportamiento. También se hacen visibles en edificios como iglesia, mezquita y sinagoga, así como en fiestas e ideas sobre la vida y la muerte. Las personas pueden elegir libremente su fe o no pertenecer a ninguna religión. Deben respetarse otras convicciones y tratarse a todas las personas con justicia.", keep_together=True)

    h1(doc, "3. Preguntas de comprensión del texto básico")
    qblock(doc, [
        ("¿Qué son las religiones?", "Son cosmovisiones y tradiciones vividas. A menudo incluyen la creencia en un poder superior o en lo sagrado y orientan valores, rituales y comunidad."),
        ("¿Cómo se llaman las cinco religiones del mundo?", "Cristianismo, islam, judaísmo, hinduismo y budismo."),
        ("¿Qué puede incluir la fe de una religión?", "Ideas sobre Dios o lo sagrado, valores y conducta, normas, costumbres, fiestas, oración, comunidad y preguntas sobre la vida y la muerte."),
        ("Nombra un edificio religioso.", "Por ejemplo, iglesia, mezquita, sinagoga o templo."),
    ], answers=True)
    callout(doc, "Precisión", "El origen de la palabra 'religión' no sustituye una respuesta sobre el contenido de la fe. Se piden ideas, valores, conducta y práctica.", CREAM, GOLD)

    h1(doc, "4. Mapa del mundo: colores, símbolos y distribución")
    table(doc, ["Religión", "Color del mapa", "Símbolo del cuaderno", "Zona principal simplificada"], [["Judaísmo", "azul", "estrella de David", "Israel; presencia mundial"], ["Budismo", "naranja", "rueda del dharma", "Asia oriental y sudoriental"], ["Cristianismo", "violeta", "cruz", "Europa, América, partes de África y Australia"], ["Hinduismo", "amarillo", "Om", "India y Asia meridional"], ["Islam", "verde", "media luna", "norte de África, Asia occidental y meridional"]], [1.15, 1.15, 1.25, 3.0], 8.8)
    paragraph(doc, "El mapa muestra zonas de concentración. Hoy existen comunidades de las cinco religiones en muchas regiones del mundo.", italic=True, color=MUTED, size=9)

    page_break(doc)
    h1(doc, "5. Gráfico de cien casillas sobre afiliación religiosa")
    table(doc, ["Grupo del ejercicio", "Casillas", "Porcentaje del gráfico"], [["Cristianos", "31", "31 %"], ["Musulmanes", "23", "23 %"], ["Hindúes", "15", "15 %"], ["Judíos", "1", "1 %"], ["Budistas", "7", "7 %"], ["otra/ninguna religión", "23", "23 %"], ["Total", "100", "100 %"]], [2.8, 1.65, 2.05], 9.6)
    callout(doc, "Importante", "Esta es la solución del gráfico dibujado de cien casillas. No es una estadística mundial actual fiable.", CREAM, GOLD)
    h1(doc, "6. Preguntas personales sobre sentido y fe")
    paragraph(doc, "Las respuestas personales son individuales. No existe una única solución correcta. Ejemplos posibles:")
    table(doc, ["Pregunta", "Posible respuesta personal"], [["¿Cómo surgió el mundo?", "La ciencia explica procesos; las religiones también transmiten relatos de sentido y origen."], ["¿Cuál es el sentido de la vida?", "Para mí: cuidar relaciones, aprender, ayudar y asumir responsabilidad."], ["¿Existe Dios?", "Las personas responden de manera diferente. Puedo respetar convicciones que no comparto."], ["¿Cómo deben convivir las personas?", "Con justicia, paz, ayuda mutua y respeto."], ["¿Qué ocurre después de la muerte?", "Las religiones ofrecen respuestas distintas; no podemos saberlo con certeza."]], [2.0, 4.5], 9.1)

    page_break(doc)
    h1(doc, "7. Comparación de las cinco religiones")
    table(doc, ["Religión", "Distribución", "Antigüedad", "Dios / personas", "Lugar", "Texto", "Normas / costumbres"], [["Cristianismo", "mundial", "unos 2000 años", "un Dios; Jesús", "iglesia", "Biblia", "amor al prójimo, fiestas"], ["Islam", "mundial", "unos 1400 años", "Alá; Mahoma", "mezquita", "Corán", "cinco pilares, Ramadán"], ["Judaísmo", "mundial", "más de 3000 años", "un Dios; Abraham/Moisés", "sinagoga", "Torá/Tanaj", "mandamientos, Shabat"], ["Budismo", "Asia; mundial", "unos 2500 años", "Buda; sin creador central", "templo", "textos", "óctuple sendero"], ["Hinduismo", "India; mundial", "más de 3000 años", "concepciones diversas", "templo", "p. ej. Vedas", "puja, fiestas"]], [.85, .85, .75, 1.25, .7, .85, 1.25], 7.4)
    paragraph(doc, "Las antigüedades son aproximadas y sirven solo para orientarse.", italic=True, color=MUTED, size=8.5)
    h1(doc, "8. Tareas sobre Jesús y el cristianismo")
    table(doc, ["Tarea", "Solución breve"], [["¿Quién fue Jesús?", "Judío de Nazaret y predicador. Para la fe cristiana, Hijo de Dios/Dios hecho ser humano."], ["Mensaje", "Amor de Dios y del prójimo: 'Ama a tu prójimo como a ti mismo'."], ["Muerte", "Crucifixión en Jerusalén bajo dominio romano."], ["Después", "Fe cristiana: resurrección al tercer día, apariciones y ascensión."], ["Fiestas", "Navidad, Viernes Santo, Pascua y Ascensión."]], [1.4, 5.1], 9.3)
    callout(doc, "Formulación precisa", "Históricamente, Jesús procedía de Nazaret. Según los evangelios nació en Belén. El material simplifica; para el examen se debe seguir la formulación utilizada en clase.", CREAM, GOLD)

    page_break(doc)
    h1(doc, "9. Tareas sobre Alá, Mahoma y el islam")
    h2(doc, "Alá")
    paragraph(doc, "Alá es la palabra árabe para Dios. En el islam hay un solo Dios. Los 99 nombres describen cualidades de Dios, por ejemplo, el Misericordioso o el Sabio.")
    h2(doc, "Mahoma: respuestas del texto con huecos y crucigrama")
    table(doc, ["N.º", "Respuesta alemana", "Relación"], [["1", "Gepriesene", "significado del nombre en el ejercicio"], ["2", "Arabien", "procedencia"], ["3", "Onkel", "crianza tras la muerte del abuelo"], ["4", "Predigten", "sermones cristianos"], ["5", "Höhle", "Hira, cerca de La Meca"], ["6", "Medina", "lugar de la emigración"], ["7", "Tod", "después de la muerte de su esposa"], ["Palabra final", "PROPHET", "mensajero que transmite el mensaje de Dios"]], [.8, 1.75, 3.95], 9.1)
    callout(doc, "Error en la pregunta", "La cueva de Hira está cerca de La Meca. Si la tarea dice 'cerca de Medina', probablemente se trata de un error de impresión.", CREAM, GOLD)
    h2(doc, "Respuesta modelo: ¿qué es un profeta?")
    paragraph(doc, "Un profeta es una persona que, según la tradición religiosa, recibe y transmite el mensaje de Dios. En el islam Mahoma es considerado el último profeta.")

    page_break(doc)
    h1(doc, "10. Etiquetas necesarias para la primera prueba")
    h2(doc, "Iglesia católica")
    table(doc, ["Término alemán", "Explicación"], [["Altar", "mesa de la eucaristía"], ["Ambo", "ambón o atril de lecturas"], ["Tabernakel", "sagrario; conserva las hostias consagradas"], ["Taufbecken", "pila bautismal"], ["Kruzifix", "cruz con representación de Jesús"], ["Orgel", "órgano para la música del culto"], ["Weihwasser", "agua bendita en la entrada; recuerda el bautismo"]], [1.5, 5.0], 9.4)
    h2(doc, "Mezquita")
    table(doc, ["Término", "Explicación"], [["Mihrab", "nicho que señala la qibla"], ["Qibla", "dirección de la oración hacia La Meca"], ["Minbar", "púlpito"], ["Gebetsraum", "sala para la oración comunitaria"], ["Gebetsteppiche", "alfombras y filas de oración"], ["Waschung", "purificación ritual"], ["Minarett", "minarete; no todas las mezquitas tienen uno"]], [1.5, 5.0], 9.4)
    h2(doc, "Sinagoga - panorama útil")
    table(doc, ["Término", "Explicación"], [["Toraschrein", "arca donde se guardan los rollos de la Torá"], ["Bima", "plataforma de lectura"], ["Ewiges Licht", "luz cercana al arca de la Torá"]], [1.5, 5.0], 9.4)

    page_break(doc)
    h1(doc, "11. Control rápido para familias")
    qblock(doc, [("¿Cuáles son las cinco religiones?", "Cristianismo, islam, judaísmo, hinduismo y budismo."), ("¿Cuáles se estudian en detalle?", "Cristianismo, islam y judaísmo."), ("¿Cuáles son monoteístas?", "Judaísmo, cristianismo e islam."), ("Cristianismo en tres palabras", "Kirche - Bibel - Jesus."), ("Islam en tres palabras", "Moschee - Koran - Mohammed."), ("Judaísmo en tres palabras", "Synagoge - Tora - Mose."), ("¿Hira?", "Cueva cerca de La Meca."), ("¿Pascua?", "Resurrección de Jesús según la fe cristiana.")], answers=True)
    h2(doc, "Fuentes")
    paragraph(doc, "Base: los materiales de clase facilitados. Revisión curricular: Lehrplan 21 del cantón de Zug, NMG.12.1, NMG.12.2 y NMG.12.5 (zg.lehrplan.ch). Las soluciones conservan el vocabulario alemán del examen y señalan simplificaciones importantes.")
    path = OUT / "Soluciones_Prueba_1_ES.docx"
    doc.save(path)
    return path


def build_solutions_english():
    doc = Document()
    configure(doc, "Solutions · NMG Religions · Test 1")
    cover(doc, "Solutions · Test 1", "World religions - Part 1", "Complete, self-contained answers and explanations", "For families and learning supporters", "English · German exam terms retained", "- Answer guide for the first test -")
    h1(doc, "How to use this guide")
    callout(doc, "Scope", "This guide answers the class tasks that are relevant to Test 1 or to the shared overview. Every heading describes the task in words and all essential information appears in this document. No photographs or file names are required. Open personal questions do not have one single correct answer.")
    bullets(doc, ["Full sentences are best for explanation questions.", "Three points are clarified: Hira is near Mecca; symbol cards simplify; the hundred-square chart is an exercise, not a current world statistic."])

    page_break(doc)
    h1(doc, "1. Learning goals for Test 1")
    h2(doc, "Test 1 content")
    bullets(doc, ["Shared overview: five world religions, broad distribution, Monotheismus/Polytheismus, questions and class content.", "Christianity: features, Jesus and a Catholic church.", "Islam: features, Mohammed and a mosque.", "Judaism: features, rules and customs."])
    callout(doc, "Not studied in detail", "The separate Buddhism and Hinduism blocks belong to Test 2. Their names, symbols and broad distribution may still appear in the shared overview.")
    h2(doc, "Model answer: monotheism and polytheism")
    paragraph(doc, "Monotheismus means belief in one God. Judaism, Christianity and Islam are monotheistic. Polytheismus means belief in several deities. Simple labels are only partly suitable for Hinduism and Buddhism.")

    h1(doc, "2. Gap-fill text: basic knowledge")
    paragraph(doc, "German answer words in gap order:", bold=True, color=DEEP)
    table(doc, ["1-8", "9-15", "16-22"], [["lateinischen\nhöhere Macht\nwissenschaftlich\nHalt\nWeltreligionen\nbesonders stark\neinen\nPolytheismus", "Wertvorstellungen\nCharakter\nVerhalten\nBauten\nKirche\nMoschee\nSynagoge", "Feiertagen\nTod\nunterschiedliche\nfrei wählen\ngar keine\nrespektieren\nbehandeln"]], [2.15, 2.15, 2.2], 9.5)
    h2(doc, "Complete meaning without gaps")
    paragraph(doc, "The word religion has a Latin origin. Many religions refer to a higher power or to the sacred; faith statements cannot be proved like scientific statements. Religion can give people support. The five world religions in this unit are Christianity, Islam, Judaism, Hinduism and Buddhism. Monotheistic religions believe in one God; polytheism means belief in several deities. Religions shape values, character and behaviour. They are also visible in buildings such as churches, mosques and synagogues, in festivals and in ideas about life and death. People may freely choose their faith or belong to no religion. Other convictions should be respected and people should be treated fairly.", keep_together=True)

    page_break(doc)
    h1(doc, "3. Comprehension questions about the basic text")
    qblock(doc, [("What are religions?", "They are worldviews and lived traditions. They often include belief in a higher power or the sacred and shape values, rituals and community."), ("What are the five world religions in this unit?", "Christianity, Islam, Judaism, Hinduism and Buddhism."), ("What can religious belief include?", "Ideas about God or the sacred, values and behaviour, rules, customs, festivals, prayer, community, life and death."), ("Name one religious building.", "For example, a church, mosque, synagogue or temple.")], answers=True)
    callout(doc, "Clarification", "The origin of the word 'religion' does not answer a question about the content of belief. The expected answer refers to ideas, values, behaviour and practice.", CREAM, GOLD)
    h1(doc, "4. World map: colours, symbols and distribution")
    table(doc, ["Religion", "Map colour", "Workbook symbol", "Simplified main area"], [["Judaism", "blue", "Star of David", "Israel; worldwide"], ["Buddhism", "orange", "Dharma wheel", "East and Southeast Asia"], ["Christianity", "violet", "cross", "Europe, the Americas, parts of Africa, Australia"], ["Hinduism", "yellow", "Om", "India and South Asia"], ["Islam", "green", "crescent", "North Africa, West and South Asia"]], [1.15, 1.15, 1.25, 3.0], 8.8)
    paragraph(doc, "The map shows concentrations. Communities of all five religions live in many world regions today.", italic=True, color=MUTED, size=9)

    page_break(doc)
    h1(doc, "5. Hundred-square chart of religious affiliation")
    table(doc, ["Group in the exercise", "Squares", "Chart percentage"], [["Christians", "31", "31 %"], ["Muslims", "23", "23 %"], ["Hindus", "15", "15 %"], ["Jews", "1", "1 %"], ["Buddhists", "7", "7 %"], ["other/no religion", "23", "23 %"], ["Total", "100", "100 %"]], [2.8, 1.65, 2.05], 9.6)
    callout(doc, "Important", "This is the answer to the drawn hundred-square chart. It is not a reliable current world-population statistic.", CREAM, GOLD)
    h1(doc, "6. Personal questions about meaning and faith")
    paragraph(doc, "Personal answers differ. There is no single correct solution. Possible model responses:")
    table(doc, ["Question", "Possible personal response"], [["How did the world begin?", "Science explains processes; religions also pass on stories about meaning and origins."], ["What is the meaning of life?", "For me: care for relationships, learn, help and take responsibility."], ["Does God exist?", "People answer differently. I can respect convictions that I do not share."], ["How should people live together?", "Fairly, peacefully, helpfully and respectfully."], ["What happens after death?", "Religions give different answers; we cannot know with certainty."]], [2.0, 4.5], 9.1)

    page_break(doc)
    h1(doc, "7. Comparison of the five religions")
    table(doc, ["Religion", "Distribution", "Age", "God / people", "Place", "Text", "Rules / customs"], [["Christianity", "worldwide", "about 2,000 years", "one God; Jesus", "church", "Bible", "love of neighbour, festivals"], ["Islam", "worldwide", "about 1,400 years", "Allah; Mohammed", "mosque", "Quran", "five pillars, Ramadan"], ["Judaism", "worldwide", "over 3,000 years", "one God; Abraham/Moses", "synagogue", "Torah/Tanakh", "commandments, Shabbat"], ["Buddhism", "Asia; worldwide", "about 2,500 years", "Buddha; no central creator", "temple", "teachings", "Eightfold Path"], ["Hinduism", "India; worldwide", "over 3,000 years", "diverse concepts", "temple", "e.g. Vedas", "puja, festivals"]], [.85, .85, .75, 1.25, .7, .85, 1.25], 7.4)
    paragraph(doc, "Ages are rounded and are provided only as orientation.", italic=True, color=MUTED, size=8.5)
    h1(doc, "8. Tasks about Jesus and Christianity")
    table(doc, ["Task", "Short answer"], [["Who was Jesus?", "A Jew from Nazareth and a preacher. In Christian belief, the Son of God/God as a human being."], ["Message", "God's love and love of neighbour: 'Love your neighbour as yourself.'"], ["Death", "Crucified in Jerusalem under Roman rule."], ["Afterwards", "Christian belief: resurrection on the third day, appearances and ascension."], ["Festivals", "Christmas, Good Friday, Easter and Ascension."]], [1.4, 5.1], 9.3)
    callout(doc, "Precise wording", "Historically, Jesus came from Nazareth. According to the Gospels, he was born in Bethlehem. The class sheet simplifies; follow the wording used in class for the test.", CREAM, GOLD)

    page_break(doc)
    h1(doc, "9. Tasks about Allah, Mohammed and Islam")
    h2(doc, "Allah")
    paragraph(doc, "Allah is the Arabic word for God. Islam teaches belief in one God. The 99 names describe qualities of God, such as the Merciful or the All-Knowing.")
    h2(doc, "Mohammed: gap-fill and crossword answers")
    table(doc, ["No.", "German answer", "Connection"], [["1", "Gepriesene", "meaning of the name in the exercise"], ["2", "Arabien", "origin"], ["3", "Onkel", "care after his grandfather's death"], ["4", "Predigten", "Christian sermons"], ["5", "Höhle", "Hira near Mecca"], ["6", "Medina", "destination of the migration"], ["7", "Tod", "after his wife's death"], ["Final word", "PROPHET", "messenger who passes on God's message"]], [.8, 1.75, 3.95], 9.1)
    callout(doc, "Error in the question", "The cave of Hira is near Mecca. If the task says 'near Medina', this is very likely a printing error.", CREAM, GOLD)
    h2(doc, "Model answer: what is a prophet?")
    paragraph(doc, "A prophet is a person who, according to religious tradition, receives and passes on God's message. In Islam, Mohammed is regarded as the final prophet.")

    page_break(doc)
    h1(doc, "10. Labels needed for Test 1")
    h2(doc, "Catholic church")
    table(doc, ["German term", "Explanation"], [["Altar", "table used for the Eucharist"], ["Ambo", "lectern for readings"], ["Tabernakel", "place where consecrated hosts are kept"], ["Taufbecken", "baptismal font"], ["Kruzifix", "cross showing Jesus"], ["Orgel", "organ used in worship and singing"], ["Weihwasser", "holy water near the entrance; recalls baptism"]], [1.5, 5.0], 9.4)
    h2(doc, "Mosque")
    table(doc, ["Term", "Explanation"], [["Mihrab", "prayer niche showing the qibla"], ["Qibla", "direction of prayer towards Mecca"], ["Minbar", "pulpit"], ["Gebetsraum", "space for communal prayer"], ["Gebetsteppiche", "prayer places and rows"], ["Waschung", "ritual washing"], ["Minarett", "minaret; not every mosque has one"]], [1.5, 5.0], 9.4)
    h2(doc, "Synagogue - useful overview")
    table(doc, ["Term", "Explanation"], [["Toraschrein", "ark where Torah scrolls are kept"], ["Bima", "raised reading platform"], ["Ewiges Licht", "light near the Torah ark"]], [1.5, 5.0], 9.4)

    page_break(doc)
    h1(doc, "11. Quick family check")
    qblock(doc, [("Which five religions?", "Christianity, Islam, Judaism, Hinduism and Buddhism."), ("Which three in detail?", "Christianity, Islam and Judaism."), ("Which three are monotheistic?", "Judaism, Christianity and Islam."), ("Christianity in three words", "Kirche - Bibel - Jesus."), ("Islam in three words", "Moschee - Koran - Mohammed."), ("Judaism in three words", "Synagoge - Tora - Mose."), ("Hira?", "A cave near Mecca."), ("Easter?", "The resurrection of Jesus according to Christian belief.")], answers=True)
    h2(doc, "Sources")
    paragraph(doc, "Based on the supplied class materials. Curriculum cross-check: Lehrplan 21 for the Canton of Zug, NMG.12.1, NMG.12.2 and NMG.12.5 (zg.lehrplan.ch). The guide retains the German exam vocabulary and flags important simplifications.")
    path = OUT / "Solutions_Test_1_EN.docx"
    doc.save(path)
    return path


def build_family_german():
    doc = Document()
    configure(doc, "Praxisleitfaden für Familien · NMG Religionen")
    cover(doc, "Familienbegleitung · Prüfung 1", "Gemeinsam sicher lernen", "Praktischer Leitfaden für Eltern und Lernbegleitende", "Für Familien von Kindern in der 6. Klasse", "Deutsch · Schweizer Schreibweise")
    h1(doc, "1. Das Ziel")
    callout(doc, "Weniger ist mehr", "Kurze, regelmässige Einheiten wirken besser als langes Auswendiglernen. Das Kind soll erklären, vergleichen, beschriften und eigene Fehler nach einer Pause nochmals lösen.")
    h2(doc, "Was für Prüfung Teil 1 wichtig ist")
    bullets(doc, ["Überblick über fünf Weltreligionen: Namen, Symbole und grobe Verbreitung.", "Monotheismus und Polytheismus verständlich erklären.", "Christentum, Islam und Judentum im Detail vergleichen.", "Stationen aus dem Leben Jesu und Mohammeds erzählen.", "Kirche, Moschee und Synagoge mit wichtigen Begriffen erkennen.", "Historische Aussagen und Glaubensaussagen respektvoll unterscheiden."])
    h2(doc, "Die Rolle der Erwachsenen")
    bullets(doc, ["Fragen stellen, statt lange zu erklären.", "Nur eine falsche Idee auf einmal korrigieren.", "Das Kind die verbesserte Antwort in eigenen Worten wiederholen lassen.", "Mit einem sichtbaren Erfolg abschliessen."])

    page_break(doc)
    h1(doc, "2. Die 15-Minuten-Routine")
    table(doc, ["Minuten", "Aktivität", "Beispiel"], [["0–3", "Abrufen ohne Unterlagen", "Nenne alles, was du über den Islam noch weisst."], ["3–7", "Eine Lernkarte klären", "Was ist der Unterschied zwischen Kirche und Synagoge?"], ["7–12", "App oder mündliche Fragen", "5–8 kurze Aufgaben beantworten."], ["12–14", "Einen Fehler verbessern", "Aus dem Fehler einen Merksatz machen."], ["14–15", "Erfolg festhalten", "Heute konnte ich ..." ]], [.8, 2.3, 3.4], 9.3)
    h2(doc, "Gute Rückfragen")
    bullets(doc, ["Woran erkennst du das?", "Was ist ähnlich, was ist verschieden?", "Welche drei Wörter gehören zusammen?", "Ist das eine historische Aussage oder eine Glaubensaussage?", "Kannst du es in zwei Sätzen erklären?", "Welche deutsche Prüfungsformulierung brauchst du?"])
    callout(doc, "Merksatz", "Kirche - Bibel - Jesus · Moschee - Koran - Mohammed · Synagoge - Tora - Mose", CREAM, GOLD)

    page_break(doc)
    h1(doc, "3. Fünf-Tage-Plan")
    table(doc, ["Tag", "Schwerpunkt", "Kurzer Lerncheck"], [["1", "Fünf Religionen, Symbole, Karte, mono/poly", "Eine Vergleichstabelle ohne Hilfe ausfüllen."], ["2", "Christentum, Jesus und Kirche", "Vier Stationen erzählen und vier Teile benennen."], ["3", "Islam, Mohammed und Moschee", "Vier Stationen und die fünf Säulen erklären."], ["4", "Judentum und Vergleich", "Drei Gemeinsamkeiten und drei Unterschiede nennen."], ["5", "Probeprüfung", "Nur Fehler wiederholen; danach ein kurzer App-Test."]], [.55, 2.8, 3.15], 9.2)
    h2(doc, "So korrigieren Sie hilfreich")
    numbers(doc, ["Zuerst benennen, was bereits stimmt.", "Eine einzige Stelle präzisieren.", "Das Kind die vollständige Antwort nochmals sagen lassen.", "Die Frage später am selben oder nächsten Tag erneut stellen."])
    h2(doc, "Wenn Motivation fehlt")
    bullets(doc, ["Mit drei leichten Fragen beginnen.", "Zwischen Sprechen, Zeichnen, Zuordnen und App wechseln.", "Ein konkretes Tagesziel vereinbaren, zum Beispiel zehn Fragen.", "Fortschritt loben, nicht nur richtige Antworten."])

    page_break(doc)
    h1(doc, "4. Sensibel über Religion sprechen")
    callout(doc, "Respektvolle Sprache", "Religionen sind vielfältig. Lernkarten, Karten und Symbole helfen beim Ordnen, beschreiben aber nie alle Menschen einer Tradition vollständig.")
    h2(doc, "Hilfreiche Formulierungen")
    bullets(doc, ["Christinnen und Christen glauben, dass ...", "Nach islamischer Überlieferung ...", "Im Judentum gibt es verschiedene Formen von ...", "Historisch lässt sich sagen ...", "Menschen können diese Frage unterschiedlich beantworten."])
    h2(doc, "Datenschutz und Lernplattform")
    paragraph(doc, "Nach der Verbindung mit dem Familienbereich werden App, Fach, Thema, Ergebnis, Punkte und Zeitpunkt gespeichert. Antworten, Fotos, Schulname und Adresse werden nicht gespeichert. Neue Themen und Fächer können später dasselbe Familienprofil verwenden.")
    h2(doc, "Fachliche Orientierung")
    paragraph(doc, "Lehrplan 21 Kanton Zug, besonders NMG.12.1, NMG.12.2 und NMG.12.5. Das Lernzielblatt der Klasse bestimmt den konkreten Prüfungsumfang. Offizielle Seiten: zg.ch/de/bildung/schulen/gemeindliche-schulen/unterricht/lehrplan21 · zg.lehrplan.ch")
    path = OUT / "Praxisleitfaden_Familie_Pruefung_1_DE.docx"
    doc.save(path)
    return path


def build_family_english():
    doc = Document()
    configure(doc, "Practical family guide · NMG Religions")
    cover(doc, "Family support · Test 1", "Learning with confidence", "A practical guide for parents and learning supporters", "For families of children in Grade 6", "English · key German exam terms retained", "—  Family learning guide for the first test  —")
    h1(doc, "1. The goal")
    callout(doc, "Short and regular", "Brief, repeated sessions work better than a long memorisation session. The learner should explain, compare, label and return to mistakes after a delay.")
    h2(doc, "What matters for the first test")
    bullets(doc, ["Overview of five world religions: names, symbols and broad geographical distribution.", "Explain Monotheismus and Polytheismus clearly.", "Compare Christianity, Islam and Judaism in detail.", "Retell key stages in the lives of Jesus and Mohammed.", "Recognise important features of a Kirche, Moschee and Synagoge.", "Distinguish historical statements from faith statements respectfully."])
    h2(doc, "The adult's role")
    bullets(doc, ["Ask questions instead of giving long explanations.", "Correct only one mistaken idea at a time.", "Ask the learner to repeat the improved answer in their own words.", "Finish with a visible success."])

    page_break(doc)
    h1(doc, "2. The 15-minute routine")
    table(doc, ["Minutes", "Activity", "Example"], [["0–3", "Recall without notes", "Tell me everything you remember about Islam."], ["3–7", "Clarify one study card", "How is a Kirche different from a Synagoge?"], ["7–12", "App or oral questions", "Answer five to eight short questions."], ["12–14", "Repair one mistake", "Turn the correction into a memory sentence."], ["14–15", "Record a success", "Today I was able to ..."]], [.8, 2.3, 3.4], 9.3)
    h2(doc, "Useful follow-up questions")
    bullets(doc, ["How do you recognise that?", "What is similar and what is different?", "Which three terms belong together?", "Is that a historical statement or a faith statement?", "Can you explain it in two sentences?", "Which German term will you need in the test?"])
    callout(doc, "Memory line", "Kirche - Bibel - Jesus · Moschee - Koran - Mohammed · Synagoge - Tora - Mose", CREAM, GOLD)

    page_break(doc)
    h1(doc, "3. Five-day plan")
    table(doc, ["Day", "Focus", "Quick check"], [["1", "Five religions, symbols, map, mono/poly", "Complete a comparison table without help."], ["2", "Christianity, Jesus and the church", "Retell four stages and label four features."], ["3", "Islam, Mohammed and the mosque", "Retell four stages and explain the five pillars."], ["4", "Judaism and comparison", "Give three similarities and three differences."], ["5", "Practice test", "Repeat mistakes only, then take a short app quiz."]], [.55, 2.8, 3.15], 9.2)
    h2(doc, "How to correct constructively")
    numbers(doc, ["Begin by naming what is already correct.", "Clarify one point only.", "Ask the learner to give the whole answer again.", "Ask the same question later that day or the next day."])
    h2(doc, "When motivation is low")
    bullets(doc, ["Start with three easy questions.", "Alternate speaking, drawing, matching and app practice.", "Agree on one concrete target, such as ten questions.", "Praise progress and persistence, not only correct answers."])

    page_break(doc)
    h1(doc, "4. Talking about religion sensitively")
    callout(doc, "Respectful language", "Religious traditions are diverse. Tables, maps and symbols are useful study aids, but they never describe every person or community completely.")
    h2(doc, "Helpful sentence starters")
    bullets(doc, ["Christians believe that ...", "According to Islamic tradition ...", "There are different Jewish approaches to ...", "Historically, we can say ...", "People may answer this question in different ways."])
    h2(doc, "Privacy and the learning platform")
    paragraph(doc, "After the learner's app is connected to the family dashboard, the platform stores the app, subject, topic, result, points and time. It does not store written answers, photos, school name or address. Future subjects can use the same family profile.")
    h2(doc, "Curriculum reference")
    paragraph(doc, "Lehrplan 21 for the Canton of Zug, especially NMG.12.1, NMG.12.2 and NMG.12.5. The class learning-objective sheet defines the exact scope of the test. Official pages: zg.ch/de/bildung/schulen/gemeindliche-schulen/unterricht/lehrplan21 · zg.lehrplan.ch")
    path = OUT / "Practical_Family_Guide_Test_1_EN.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for result in (
        build_german(),
        build_spanish(),
        build_solutions_german(),
        build_solutions_spanish(),
        build_solutions_english(),
        build_family_german(),
        build_family_english(),
    ):
        print(result)
